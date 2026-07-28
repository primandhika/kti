#!/usr/bin/env python3
"""
compile_docx.py — Kompilasi naskah buku_presentasi_ilmiah dari .md menjadi .docx

Diadaptasi dari script buku_media agar dapat dipakai untuk proyek buku
Mahir Presentasi Ilmiah dengan Strategi Microlearning dan Metakognitif.

Dependensi:
    pip install python-docx
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx belum terinstal.")
    print("Jalankan:  pip install python-docx")
    sys.exit(1)

BASE_DIR = Path(__file__).resolve().parent
BAB_DIR = BASE_DIR / "bab"
COVER_FILES = {"COVER.md"}
FONT_NAME = "Times New Roman"

PAGE_PRESETS = {
    "A5": {
        "page_width": 14.8,
        "page_height": 21.0,
        "margin_top": 1.5,
        "margin_bottom": 1.5,
        "margin_dalam": 2.0,
        "margin_luar": 1.5,
        "font_body": 11,
        "font_h1": 14,
        "font_h2": 12,
        "font_h3": 11,
        "font_h4": 11,
        "font_table": 9,
        "font_code": 8,
        "font_caption": 9,
        "font_cover_title": 18,
        "font_cover_subtitle": 13,
        "font_cover_author": 12,
        "line_spacing": 1.3,
        "first_indent": 1.0,
        "list_indent": 0.8,
        "code_indent": 0.6,
    },
    "B5": {
        "page_width": 17.6,
        "page_height": 25.0,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_dalam": 2.5,
        "margin_luar": 1.8,
        "font_body": 11,
        "font_h1": 15,
        "font_h2": 13,
        "font_h3": 12,
        "font_h4": 11,
        "font_table": 9,
        "font_code": 8.5,
        "font_caption": 9,
        "font_cover_title": 20,
        "font_cover_subtitle": 14,
        "font_cover_author": 13,
        "line_spacing": 1.4,
        "first_indent": 1.0,
        "list_indent": 0.8,
        "code_indent": 0.8,
    },
    "A4": {
        "page_width": 21.0,
        "page_height": 29.7,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_dalam": 3.0,
        "margin_luar": 2.54,
        "font_body": 12,
        "font_h1": 16,
        "font_h2": 14,
        "font_h3": 13,
        "font_h4": 12,
        "font_table": 10,
        "font_code": 9,
        "font_caption": 10,
        "font_cover_title": 22,
        "font_cover_subtitle": 16,
        "font_cover_author": 14,
        "line_spacing": 1.5,
        "first_indent": 1.27,
        "list_indent": 1.0,
        "code_indent": 1.0,
    },
}


def discover_book_files() -> list[Path]:
    """Bangun urutan file buku secara fleksibel."""
    awal = [
        BASE_DIR / "COVER.md",
        BASE_DIR / "HAK_CIPTA.md",
        BASE_DIR / "PENGANTAR.md",
        BASE_DIR / "DAFTAR_ISI.md",
        BASE_DIR / "DAFTAR_GAMBAR.md",
        BASE_DIR / "DAFTAR_TABEL.md",
        BASE_DIR / "CARA_MENGGUNAKAN_BUKU.md",
    ]
    bab = sorted(BAB_DIR.glob("*.md")) if BAB_DIR.exists() else []
    akhir = [
        BASE_DIR / "GLOSARIUM.md",
        BASE_DIR / "DAFTAR_PUSTAKA.md",
        BASE_DIR / "INDEKS.md",
        BASE_DIR / "TENTANG_PENULIS.md",
    ]
    return awal + bab + akhir


def get_config(args) -> dict:
    preset_name = args.ukuran.upper()
    if preset_name in PAGE_PRESETS:
        cfg = dict(PAGE_PRESETS[preset_name])
    else:
        if not args.lebar or not args.tinggi:
            print(f"ERROR: Ukuran '{args.ukuran}' tidak dikenali.")
            print(f"Preset tersedia: {', '.join(PAGE_PRESETS.keys())}")
            print("Atau gunakan --lebar dan --tinggi untuk ukuran kustom.")
            sys.exit(1)
        cfg = dict(PAGE_PRESETS["A5"])
        cfg["page_width"] = args.lebar
        cfg["page_height"] = args.tinggi

    if args.margin_top is not None:
        cfg["margin_top"] = args.margin_top
    if args.margin_bottom is not None:
        cfg["margin_bottom"] = args.margin_bottom
    if args.margin_dalam is not None:
        cfg["margin_dalam"] = args.margin_dalam
    if args.margin_luar is not None:
        cfg["margin_luar"] = args.margin_luar
    if args.font_body is not None:
        cfg["font_body"] = args.font_body
    if args.spasi is not None:
        cfg["line_spacing"] = args.spasi

    cfg["mirror"] = args.mirror
    return cfg


class MarkdownParser:
    def __init__(self, doc: Document, cfg: dict):
        self.doc = doc
        self.cfg = cfg

    @staticmethod
    def _clean_html_br(text: str) -> str:
        return re.sub(r"<br\s*/?>", "\n", text)

    def _add_inline_formatting(self, paragraph, text: str):
        text = self._clean_html_br(text)
        pattern = re.compile(
            r"(\*\*\*(.+?)\*\*\*)"
            r"|(\*\*(.+?)\*\*)"
            r"|(\*(.+?)\*)"
            r"|(_(.+?)_)"
            r"|(`(.+?)`)"
            r"|([^*_`]+)"
        )
        for m in pattern.finditer(text):
            if m.group(2):
                run = paragraph.add_run(m.group(2))
                run.bold = True
                run.italic = True
            elif m.group(4):
                run = paragraph.add_run(m.group(4))
                run.bold = True
            elif m.group(6):
                run = paragraph.add_run(m.group(6))
                run.italic = True
            elif m.group(8):
                run = paragraph.add_run(m.group(8))
                run.italic = True
            elif m.group(10):
                run = paragraph.add_run(m.group(10))
                run.font.name = "Consolas"
                run.font.size = Pt(self.cfg["font_code"])
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif m.group(11):
                paragraph.add_run(m.group(11))

    def _format_paragraph(self, paragraph, font_size=None,
                          bold=False, italic=False,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=6,
                          first_line_indent=None):
        if font_size is None:
            font_size = self.cfg["font_body"]

        pf = paragraph.paragraph_format
        pf.alignment = alignment
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = self.cfg["line_spacing"]

        if first_line_indent is not None:
            pf.first_line_indent = Cm(first_line_indent)

        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(font_size)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'
                )
                rpr.append(rFonts)
            else:
                rFonts.set(qn("w:eastAsia"), FONT_NAME)
            if bold:
                run.bold = True
            if italic:
                run.italic = True

    def _parse_table(self, lines: list[str]) -> list[list[str]]:
        rows = []
        for line in lines:
            line = line.strip()
            if not line.startswith("|"):
                continue
            if re.match(r"^\|[\s\-:|]+\|$", line):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
        return rows

    def _add_table(self, rows: list[list[str]]):
        if not rows:
            return
        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        ft = self.cfg["font_table"]
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.paragraphs[0].clear()
                    self._add_inline_formatting(cell.paragraphs[0], cell_text)
                    for run in cell.paragraphs[0].runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(ft)
                        if i == 0:
                            run.bold = True
                    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.15
                    if i == 0:
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
                        )
                        cell._element.get_or_add_tcPr().append(shading)
        self.doc.add_paragraph()

    def _is_image_placeholder(self, line: str) -> bool:
        stripped = line.strip()
        return (stripped.startswith("[") and stripped.endswith("]")
                and not stripped.startswith("[!"))

    def _add_image_placeholder(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(self.cfg["font_caption"])
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)

    def _add_image_caption(self, text: str):
        p = self.doc.add_paragraph()
        self._add_inline_formatting(p, text)
        self._format_paragraph(
            p,
            font_size=self.cfg["font_caption"],
            italic=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=2,
            space_after=12,
        )

    def _add_heading(self, text: str, level: int, is_cover: bool = False):
        size_map = {
            1: self.cfg["font_h1"],
            2: self.cfg["font_h2"],
            3: self.cfg["font_h3"],
            4: self.cfg["font_h4"],
        }
        font_size = size_map.get(level, self.cfg["font_body"])
        p = self.doc.add_paragraph()
        self._add_inline_formatting(p, text)
        alignment = (WD_ALIGN_PARAGRAPH.CENTER
                     if (level <= 2 or is_cover)
                     else WD_ALIGN_PARAGRAPH.LEFT)
        space_before = 24 if level == 1 else (18 if level == 2 else 12)
        space_after = 12 if level <= 2 else 6
        self._format_paragraph(p, font_size=font_size, bold=True,
                               alignment=alignment,
                               space_before=space_before,
                               space_after=space_after)

    def _add_list_item(self, text: str, ordered: bool = False, number: int = 1):
        p = self.doc.add_paragraph()
        prefix = f"{number}. " if ordered else "• "
        p.add_run(prefix)
        self._add_inline_formatting(p, text)
        self._format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               space_before=0, space_after=3,
                               first_line_indent=0)
        p.paragraph_format.left_indent = Cm(self.cfg["list_indent"])

    def _add_code_block(self, lines: list[str]):
        for line in lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(self.cfg["font_code"])
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Cm(self.cfg["code_indent"])

    def _render_cover(self, lines: list[str]):
        blank = 4 if self.cfg["page_height"] <= 21.5 else 6
        for _ in range(blank):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)

        for line in lines:
            line = line.rstrip()
            if not line:
                continue
            if line.startswith("# "):
                text = line[2:].strip()
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(self.cfg["font_cover_title"])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
            elif line.startswith("## "):
                text = line[3:].strip()
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                run.font.name = FONT_NAME
                run.font.size = Pt(self.cfg["font_cover_subtitle"])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(20)
            else:
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                text = re.sub(r"<br\s*/?>", "", text).strip()
                if not text:
                    continue
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = FONT_NAME
                run.font.size = Pt(self.cfg["font_cover_author"])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)

    def render_file(self, filepath: Path, is_cover: bool = False,
                    add_page_break: bool = True):
        if not filepath.exists():
            print(f"  ⚠ File tidak ditemukan: {filepath}")
            return

        text = filepath.read_text(encoding="utf-8")
        lines = text.split("\n")

        if is_cover:
            self._render_cover(lines)
            return
        if add_page_break:
            self.doc.add_page_break()

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if not stripped:
                i += 1
                continue

            if stripped.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1
                self._add_code_block(code_lines)
                continue

            if stripped.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                rows = self._parse_table(table_lines)
                self._add_table(rows)
                continue

            heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                self._add_heading(heading_match.group(2), level, is_cover=is_cover)
                i += 1
                continue

            if self._is_image_placeholder(stripped):
                self._add_image_placeholder(stripped)
                i += 1
                if i < len(lines) and lines[i].strip().startswith("Gambar"):
                    self._add_image_caption(lines[i].strip())
                    i += 1
                continue

            if stripped.startswith("Gambar ") and re.match(r"^Gambar\s+\d+\.\d+", stripped):
                self._add_image_caption(stripped)
                i += 1
                continue

            ol_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
            if ol_match:
                num = int(ol_match.group(1))
                self._add_list_item(ol_match.group(2), ordered=True, number=num)
                i += 1
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                self._add_list_item(stripped[2:], ordered=False)
                i += 1
                continue

            para_lines = [stripped]
            i += 1
            while i < len(lines):
                nl = lines[i].strip()
                if (not nl
                        or nl.startswith("#")
                        or nl.startswith("|")
                        or nl.startswith("```")
                        or nl.startswith("- ")
                        or nl.startswith("* ")
                        or re.match(r"^\d+\.\s+", nl)
                        or self._is_image_placeholder(nl)):
                    break
                para_lines.append(nl)
                i += 1

            full_text = " ".join(para_lines)
            p = self.doc.add_paragraph()
            self._add_inline_formatting(p, full_text)
            self._format_paragraph(
                p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0,
                space_after=6,
                first_line_indent=self.cfg["first_indent"],
            )


def setup_document(cfg: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(cfg["page_width"])
    section.page_height = Cm(cfg["page_height"])
    section.top_margin = Cm(cfg["margin_top"])
    section.bottom_margin = Cm(cfg["margin_bottom"])

    if cfg["mirror"]:
        section.left_margin = Cm(cfg["margin_dalam"])
        section.right_margin = Cm(cfg["margin_luar"])
        sectPr = section._sectPr
        mirrorMargins = sectPr.find(qn("w:mirrorMargins"))
        if mirrorMargins is None:
            sectPr.append(parse_xml(f'<w:mirrorMargins {nsdecls("w")}/>'))
    else:
        section.left_margin = Cm(cfg["margin_dalam"])
        section.right_margin = Cm(cfg["margin_luar"])

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(cfg["font_body"])
    style.paragraph_format.line_spacing = cfg["line_spacing"]
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rpr.append(rFonts)

    return doc


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="Kompilasi buku_presentasi_ilmiah dari .md menjadi .docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    p.add_argument(
        "--output", "-o",
        default="Buku_Mahir_Presentasi_Ilmiah.docx",
        help="Nama file output (default: Buku_Mahir_Presentasi_Ilmiah.docx)",
    )
    g_page = p.add_argument_group("ukuran halaman")
    g_page.add_argument(
        "--ukuran", "-u",
        default="A5",
        help="Preset ukuran: A5 (default), B5, A4, atau custom",
    )
    g_page.add_argument("--lebar", type=float, default=None,
                        help="Lebar halaman dalam cm untuk ukuran custom")
    g_page.add_argument("--tinggi", type=float, default=None,
                        help="Tinggi halaman dalam cm untuk ukuran custom")

    g_margin = p.add_argument_group("margin (cm)")
    g_margin.add_argument("--margin-top", type=float, default=None)
    g_margin.add_argument("--margin-bottom", type=float, default=None)
    g_margin.add_argument("--margin-dalam", type=float, default=None)
    g_margin.add_argument("--margin-luar", type=float, default=None)
    g_margin.add_argument("--mirror", action="store_true", default=False)

    g_typo = p.add_argument_group("tipografi")
    g_typo.add_argument("--font-body", type=float, default=None)
    g_typo.add_argument("--spasi", type=float, default=None)
    return p


def main():
    parser = build_parser()
    args = parser.parse_args()
    cfg = get_config(args)
    output_path = BASE_DIR / args.output
    urutan_file = discover_book_files()
    ukuran_label = args.ukuran.upper()

    print("=" * 60)
    print("  Kompilasi Buku Mahir Presentasi Ilmiah → DOCX")
    print("=" * 60)
    print(f"\n  Ukuran halaman : {ukuran_label} ({cfg['page_width']}×{cfg['page_height']} cm)")
    print(f"  Margin (cm)    : atas={cfg['margin_top']}  bawah={cfg['margin_bottom']}  dalam={cfg['margin_dalam']}  luar={cfg['margin_luar']}")
    print(f"  Mirror margin  : {'Ya' if cfg['mirror'] else 'Tidak'}")
    print(f"  Font isi       : {FONT_NAME} {cfg['font_body']}pt, spasi {cfg['line_spacing']}")
    print()

    missing = [f for f in urutan_file if not f.exists()]
    if missing:
        print("⚠ File berikut tidak ditemukan:")
        for f in missing:
            print(f"    {f.relative_to(BASE_DIR)}")
        resp = input("\nLanjutkan tanpa file tersebut? (y/n): ").strip().lower()
        if resp != "y":
            print("Dibatalkan.")
            sys.exit(0)

    doc = setup_document(cfg)
    md = MarkdownParser(doc, cfg)

    total = len(urutan_file)
    for idx, filepath in enumerate(urutan_file, 1):
        filename = filepath.name
        is_cover = filename in COVER_FILES
        add_break = (idx > 1) and not is_cover
        rel = filepath.relative_to(BASE_DIR)
        print(f"  [{idx:2d}/{total}] {rel}")
        md.render_file(filepath, is_cover=is_cover, add_page_break=add_break)

    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f"\n✅ Berhasil! File disimpan:")
    print(f"   {output_path}")
    print(f"   Ukuran: {size_kb:.0f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
