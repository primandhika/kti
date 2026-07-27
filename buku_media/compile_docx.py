#!/usr/bin/env python3
"""
compile_docx.py — Kompilasi naskah buku_media dari .md menjadi .docx

Lihat README_COMPILE.md untuk dokumentasi lengkap.

Dependensi:
    pip install python-docx

Contoh cepat:
    python compile_docx.py                          # A5 default (cetak buku)
    python compile_docx.py --ukuran A4              # A4 untuk draf/review
    python compile_docx.py --ukuran A5 --mirror     # A5 + mirror margin (jilid)
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx belum terinstal.")
    print("Jalankan:  pip install python-docx")
    sys.exit(1)

# ─────────────────────────────────────────────
# Konfigurasi path
# ─────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
BAB_DIR = BASE_DIR / "bab"

# Urutan file untuk kompilasi buku
URUTAN_FILE = [
    # Bagian awal
    BASE_DIR / "COVER.md",
    BASE_DIR / "HAK_CIPTA.md",
    BASE_DIR / "PENGANTAR.md",
    BASE_DIR / "DAFTAR_ISI.md",
    BASE_DIR / "DAFTAR_GAMBAR.md",
    BASE_DIR / "DAFTAR_TABEL.md",
    BASE_DIR / "CARA_MENGGUNAKAN_BUKU.md",
    # Bab-bab isi
    BAB_DIR / "01_Konsep_Dasar_Media_Pembelajaran.md",
    BAB_DIR / "02_Media_Pembelajaran_di_Era_Digital_dan_AI.md",
    BAB_DIR / "03_Analisis_Kebutuhan_dan_Perencanaan_Media.md",
    BAB_DIR / "04_Desain_dan_Evaluasi_Media_Pembelajaran.md",
    BAB_DIR / "05_Media_untuk_Pembelajaran_Keterampilan_Berbahasa.md",
    BAB_DIR / "06_Media_untuk_Pembelajaran_Sastra_Indonesia.md",
    BAB_DIR / "07_Studio_Kreasi_Media_Digital.md",
    BAB_DIR / "08_Studio_Media_Interaktif_Tanpa_Kode.md",
    BAB_DIR / "09_Studio_AI_dan_Vibe_Coding_untuk_Media_Pembelajaran.md",
    BAB_DIR / "10_Implementasi_Uji_Penggunaan_dan_Portofolio_Media.md",
    # Bagian akhir
    BASE_DIR / "GLOSARIUM.md",
    BASE_DIR / "DAFTAR_PUSTAKA.md",
    BASE_DIR / "INDEKS.md",
    BASE_DIR / "TENTANG_PENULIS.md",
]

# File yang mendapat perlakuan halaman sampul
COVER_FILES = {"COVER.md"}

# Font utama
FONT_NAME = "Times New Roman"

# ─────────────────────────────────────────────
# Preset ukuran halaman
# ─────────────────────────────────────────────
# Masing-masing preset berisi:
#   page_width, page_height (cm),
#   margin_top, margin_bottom, margin_dalam, margin_luar (cm),
#   font_body, font_h1, font_h2, font_h3, font_h4,
#   font_table, font_code, font_caption,
#   font_cover_title, font_cover_subtitle, font_cover_author,
#   line_spacing, first_indent (cm),
#   list_indent (cm), code_indent (cm)

PAGE_PRESETS = {
    "A5": {
        "page_width": 14.8,
        "page_height": 21.0,
        # Margin cetak buku A5 – dalam (gutter/jilid) lebih besar
        "margin_top": 1.5,
        "margin_bottom": 1.5,
        "margin_dalam": 2.0,      # sisi jilid
        "margin_luar": 1.5,       # sisi potong
        # Tipografi proporsional untuk A5
        "font_body": 11,
        "font_h1": 14,
        "font_h2": 12,
        "font_h3": 11,
        "font_h4": 11,
        "font_table": 9,
        "font_code": 8,
        "font_caption": 9,
        "font_cover_title": 18,
        "font_cover_subtitle": 13,
        "font_cover_author": 12,
        "line_spacing": 1.3,
        "first_indent": 1.0,
        "list_indent": 0.8,
        "code_indent": 0.6,
    },
    "B5": {
        "page_width": 17.6,
        "page_height": 25.0,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_dalam": 2.5,
        "margin_luar": 1.8,
        "font_body": 11,
        "font_h1": 15,
        "font_h2": 13,
        "font_h3": 12,
        "font_h4": 11,
        "font_table": 9,
        "font_code": 8.5,
        "font_caption": 9,
        "font_cover_title": 20,
        "font_cover_subtitle": 14,
        "font_cover_author": 13,
        "line_spacing": 1.4,
        "first_indent": 1.0,
        "list_indent": 0.8,
        "code_indent": 0.8,
    },
    "A4": {
        "page_width": 21.0,
        "page_height": 29.7,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_dalam": 3.0,
        "margin_luar": 2.54,
        "font_body": 12,
        "font_h1": 16,
        "font_h2": 14,
        "font_h3": 13,
        "font_h4": 12,
        "font_table": 10,
        "font_code": 9,
        "font_caption": 10,
        "font_cover_title": 22,
        "font_cover_subtitle": 16,
        "font_cover_author": 14,
        "line_spacing": 1.5,
        "first_indent": 1.27,
        "list_indent": 1.0,
        "code_indent": 1.0,
    },
}


def get_config(args) -> dict:
    """Gabungkan preset dengan override CLI menjadi konfigurasi final."""
    preset_name = args.ukuran.upper()
    if preset_name in PAGE_PRESETS:
        cfg = dict(PAGE_PRESETS[preset_name])
    else:
        # Custom: wajib --lebar dan --tinggi
        if not args.lebar or not args.tinggi:
            print(f"ERROR: Ukuran '{args.ukuran}' tidak dikenali.")
            print(f"Preset tersedia: {', '.join(PAGE_PRESETS.keys())}")
            print("Atau gunakan --lebar dan --tinggi untuk ukuran kustom.")
            sys.exit(1)
        # Mulai dari A5 sebagai basis, lalu timpa dimensi
        cfg = dict(PAGE_PRESETS["A5"])
        cfg["page_width"] = args.lebar
        cfg["page_height"] = args.tinggi

    # Override margin jika diberikan lewat CLI
    if args.margin_top is not None:
        cfg["margin_top"] = args.margin_top
    if args.margin_bottom is not None:
        cfg["margin_bottom"] = args.margin_bottom
    if args.margin_dalam is not None:
        cfg["margin_dalam"] = args.margin_dalam
    if args.margin_luar is not None:
        cfg["margin_luar"] = args.margin_luar

    # Override tipografi jika diberikan
    if args.font_body is not None:
        cfg["font_body"] = args.font_body
    if args.spasi is not None:
        cfg["line_spacing"] = args.spasi

    # Mirror margin
    cfg["mirror"] = args.mirror

    return cfg


# ─────────────────────────────────────────────
# Pemrosesan Markdown
# ─────────────────────────────────────────────

class MarkdownParser:
    """Parser Markdown sederhana yang langsung menulis ke python-docx."""

    def __init__(self, doc: Document, cfg: dict):
        self.doc = doc
        self.cfg = cfg

    # ── Utilitas teks ──

    @staticmethod
    def _clean_html_br(text: str) -> str:
        """Hapus tag <br> / <br/>."""
        return re.sub(r"<br\s*/?>", "\n", text)

    def _add_inline_formatting(self, paragraph, text: str):
        """Proses bold, italic, inline code, dan teks biasa ke runs."""
        text = self._clean_html_br(text)

        pattern = re.compile(
            r"(\*\*\*(.+?)\*\*\*)"   # bold italic
            r"|(\*\*(.+?)\*\*)"      # bold
            r"|(\*(.+?)\*)"          # italic
            r"|(_(.+?)_)"            # italic alt
            r"|(`(.+?)`)"            # inline code
            r"|([^*_`]+)"            # teks biasa
        )

        for m in pattern.finditer(text):
            if m.group(2):  # bold italic
                run = paragraph.add_run(m.group(2))
                run.bold = True
                run.italic = True
            elif m.group(4):  # bold
                run = paragraph.add_run(m.group(4))
                run.bold = True
            elif m.group(6):  # italic
                run = paragraph.add_run(m.group(6))
                run.italic = True
            elif m.group(8):  # italic alt
                run = paragraph.add_run(m.group(8))
                run.italic = True
            elif m.group(10):  # inline code
                run = paragraph.add_run(m.group(10))
                run.font.name = "Consolas"
                run.font.size = Pt(self.cfg["font_code"])
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif m.group(11):  # teks biasa
                paragraph.add_run(m.group(11))

    def _format_paragraph(self, paragraph, font_size=None,
                          bold=False, italic=False,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=6,
                          first_line_indent=None):
        """Atur format paragraf."""
        if font_size is None:
            font_size = self.cfg["font_body"]

        pf = paragraph.paragraph_format
        pf.alignment = alignment
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = self.cfg["line_spacing"]

        if first_line_indent is not None:
            pf.first_line_indent = Cm(first_line_indent)

        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(font_size)
            # Atur font eastAsia
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'
                )
                rpr.append(rFonts)
            else:
                rFonts.set(qn("w:eastAsia"), FONT_NAME)

            if bold:
                run.bold = True
            if italic:
                run.italic = True

    # ── Tabel ──

    def _parse_table(self, lines: list[str]) -> list[list[str]]:
        """Parse baris-baris tabel markdown menjadi list of rows."""
        rows = []
        for line in lines:
            line = line.strip()
            if not line.startswith("|"):
                continue
            if re.match(r"^\|[\s\-:|]+\|$", line):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
        return rows

    def _add_table(self, rows: list[list[str]]):
        """Tambahkan tabel ke dokumen."""
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        ft = self.cfg["font_table"]
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.paragraphs[0].clear()
                    self._add_inline_formatting(cell.paragraphs[0], cell_text)

                    for run in cell.paragraphs[0].runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(ft)
                        if i == 0:
                            run.bold = True

                    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.15

                    if i == 0:
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
                        )
                        cell._element.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph()

    # ── Penanda gambar ──

    def _is_image_placeholder(self, line: str) -> bool:
        """Cek apakah baris berupa penanda tempat gambar."""
        stripped = line.strip()
        return (stripped.startswith("[") and stripped.endswith("]")
                and not stripped.startswith("[!"))

    def _add_image_placeholder(self, text: str):
        """Tambahkan placeholder gambar sebagai teks abu-abu."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(self.cfg["font_caption"])
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)

    def _add_image_caption(self, text: str):
        """Tambahkan caption gambar."""
        p = self.doc.add_paragraph()
        self._add_inline_formatting(p, text)
        self._format_paragraph(p, font_size=self.cfg["font_caption"],
                               italic=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               space_before=2, space_after=12)

    # ── Heading ──

    def _add_heading(self, text: str, level: int, is_cover: bool = False):
        """Tambahkan heading ke dokumen."""
        size_map = {
            1: self.cfg["font_h1"],
            2: self.cfg["font_h2"],
            3: self.cfg["font_h3"],
            4: self.cfg["font_h4"],
        }
        font_size = size_map.get(level, self.cfg["font_body"])

        p = self.doc.add_paragraph()
        self._add_inline_formatting(p, text)
        alignment = (WD_ALIGN_PARAGRAPH.CENTER
                     if (level <= 2 or is_cover)
                     else WD_ALIGN_PARAGRAPH.LEFT)
        space_before = 24 if level == 1 else (18 if level == 2 else 12)
        space_after = 12 if level <= 2 else 6
        self._format_paragraph(p, font_size=font_size, bold=True,
                               alignment=alignment,
                               space_before=space_before,
                               space_after=space_after)

    # ── List ──

    def _add_list_item(self, text: str, ordered: bool = False, number: int = 1):
        """Tambahkan item daftar."""
        p = self.doc.add_paragraph()
        prefix = f"{number}. " if ordered else "• "
        p.add_run(prefix)
        self._add_inline_formatting(p, text)
        self._format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               space_before=0, space_after=3,
                               first_line_indent=0)
        p.paragraph_format.left_indent = Cm(self.cfg["list_indent"])

    # ── Blok kode ──

    def _add_code_block(self, lines: list[str]):
        """Tambahkan blok kode."""
        for line in lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(self.cfg["font_code"])
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Cm(self.cfg["code_indent"])

    # ── Cover ──

    def _render_cover(self, lines: list[str]):
        """Render halaman sampul dengan format centered."""
        # Jumlah baris kosong disesuaikan ukuran halaman
        blank = 4 if self.cfg["page_height"] <= 21.5 else 6
        for _ in range(blank):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)

        for line in lines:
            line = line.rstrip()
            if not line:
                continue

            if line.startswith("# "):
                text = line[2:].strip()
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.name = FONT_NAME
                run.font.size = Pt(self.cfg["font_cover_title"])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)

            elif line.startswith("## "):
                text = line[3:].strip()
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.italic = True
                run.font.name = FONT_NAME
                run.font.size = Pt(self.cfg["font_cover_subtitle"])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(20)

            else:
                text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                text = re.sub(r"<br\s*/?>", "", text).strip()
                if not text:
                    continue
                p = self.doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = FONT_NAME
                run.font.size = Pt(self.cfg["font_cover_author"])
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(4)

    # ── Parser utama per file ──

    def render_file(self, filepath: Path, is_cover: bool = False,
                    add_page_break: bool = True):
        """Parse dan render satu file markdown ke dokumen."""
        if not filepath.exists():
            print(f"  ⚠ File tidak ditemukan: {filepath}")
            return

        text = filepath.read_text(encoding="utf-8")
        lines = text.split("\n")

        if is_cover:
            self._render_cover(lines)
            return

        if add_page_break:
            self.doc.add_page_break()

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Baris kosong
            if not stripped:
                i += 1
                continue

            # Blok kode
            if stripped.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1
                self._add_code_block(code_lines)
                continue

            # Tabel
            if stripped.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                rows = self._parse_table(table_lines)
                self._add_table(rows)
                continue

            # Heading
            heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                self._add_heading(heading_match.group(2), level,
                                  is_cover=is_cover)
                i += 1
                continue

            # Placeholder gambar
            if self._is_image_placeholder(stripped):
                self._add_image_placeholder(stripped)
                i += 1
                if (i < len(lines)
                        and lines[i].strip().startswith("Gambar")):
                    self._add_image_caption(lines[i].strip())
                    i += 1
                continue

            # Caption gambar standalone
            if (stripped.startswith("Gambar ")
                    and re.match(r"^Gambar\s+\d+\.\d+", stripped)):
                self._add_image_caption(stripped)
                i += 1
                continue

            # Ordered list
            ol_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
            if ol_match:
                num = int(ol_match.group(1))
                self._add_list_item(ol_match.group(2), ordered=True,
                                    number=num)
                i += 1
                continue

            # Unordered list
            if stripped.startswith("- ") or stripped.startswith("* "):
                self._add_list_item(stripped[2:], ordered=False)
                i += 1
                continue

            # Paragraf biasa — kumpulkan sampai pemisah
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                nl = lines[i].strip()
                if (not nl
                        or nl.startswith("#")
                        or nl.startswith("|")
                        or nl.startswith("```")
                        or nl.startswith("- ")
                        or nl.startswith("* ")
                        or re.match(r"^\d+\.\s+", nl)
                        or self._is_image_placeholder(nl)):
                    break
                para_lines.append(nl)
                i += 1

            full_text = " ".join(para_lines)
            p = self.doc.add_paragraph()
            self._add_inline_formatting(p, full_text)
            self._format_paragraph(
                p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6,
                first_line_indent=self.cfg["first_indent"],
            )


# ─────────────────────────────────────────────
# Setup dokumen
# ─────────────────────────────────────────────

def setup_document(cfg: dict) -> Document:
    """Buat dokumen baru dengan pengaturan halaman sesuai konfigurasi."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(cfg["page_width"])
    section.page_height = Cm(cfg["page_height"])
    section.top_margin = Cm(cfg["margin_top"])
    section.bottom_margin = Cm(cfg["margin_bottom"])

    if cfg["mirror"]:
        # Mirror margin: gutter di sisi jilid, bergantian kiri-kanan
        section.left_margin = Cm(cfg["margin_dalam"])
        section.right_margin = Cm(cfg["margin_luar"])
        # Aktifkan mirror margins lewat XML
        sectPr = section._sectPr
        mirrorMargins = sectPr.find(qn("w:mirrorMargins"))
        if mirrorMargins is None:
            sectPr.append(parse_xml(
                f'<w:mirrorMargins {nsdecls("w")}/>'
            ))
    else:
        # Tanpa mirror: margin dalam di kiri, luar di kanan
        section.left_margin = Cm(cfg["margin_dalam"])
        section.right_margin = Cm(cfg["margin_luar"])

    # Style default
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(cfg["font_body"])
    style.paragraph_format.line_spacing = cfg["line_spacing"]
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'
        )
        rpr.append(rFonts)

    return doc


# ─────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────

def build_parser() -> argparse.ArgumentParser:
    """Bangun argument parser dengan semua opsi."""
    p = argparse.ArgumentParser(
        description="Kompilasi buku_media dari .md menjadi .docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
contoh penggunaan:
  python compile_docx.py                                    # A5 cetak buku
  python compile_docx.py --ukuran A4                        # A4 draf review
  python compile_docx.py --ukuran B5                        # B5 cetak buku
  python compile_docx.py --ukuran A5 --mirror               # A5 + mirror margin
  python compile_docx.py --ukuran A5 --margin-dalam 2.5     # override margin jilid
  python compile_docx.py --ukuran custom --lebar 15 --tinggi 23
  python compile_docx.py --font-body 10 --spasi 1.15        # override tipografi
        """,
    )

    # Output
    p.add_argument(
        "--output", "-o",
        default="Buku_Media_Pembelajaran.docx",
        help="Nama file output (default: Buku_Media_Pembelajaran.docx)",
    )

    # Ukuran halaman
    g_page = p.add_argument_group("ukuran halaman")
    g_page.add_argument(
        "--ukuran", "-u",
        default="A5",
        help="Preset ukuran: A5 (default, cetak buku), B5, A4, "
             "atau 'custom' (wajib --lebar --tinggi)",
    )
    g_page.add_argument("--lebar", type=float, default=None,
                        help="Lebar halaman dalam cm (untuk ukuran custom)")
    g_page.add_argument("--tinggi", type=float, default=None,
                        help="Tinggi halaman dalam cm (untuk ukuran custom)")

    # Margin
    g_margin = p.add_argument_group("margin (cm)")
    g_margin.add_argument("--margin-top", type=float, default=None,
                          help="Margin atas")
    g_margin.add_argument("--margin-bottom", type=float, default=None,
                          help="Margin bawah")
    g_margin.add_argument("--margin-dalam", type=float, default=None,
                          help="Margin dalam / sisi jilid (gutter)")
    g_margin.add_argument("--margin-luar", type=float, default=None,
                          help="Margin luar / sisi potong")
    g_margin.add_argument("--mirror", action="store_true", default=False,
                          help="Aktifkan mirror margin (gutter bergantian "
                               "kiri-kanan untuk cetak bolak-balik)")

    # Tipografi
    g_typo = p.add_argument_group("tipografi")
    g_typo.add_argument("--font-body", type=float, default=None,
                        help="Ukuran font isi (pt)")
    g_typo.add_argument("--spasi", type=float, default=None,
                        help="Spasi baris (misal 1.3, 1.5)")

    return p


def main():
    parser = build_parser()
    args = parser.parse_args()
    cfg = get_config(args)

    output_path = BASE_DIR / args.output
    ukuran_label = args.ukuran.upper()

    print("=" * 60)
    print("  Kompilasi Buku Media Pembelajaran → DOCX")
    print("=" * 60)

    # Tampilkan konfigurasi
    print(f"\n  Ukuran halaman : {ukuran_label} "
          f"({cfg['page_width']}×{cfg['page_height']} cm)")
    print(f"  Margin (cm)    : atas={cfg['margin_top']}  "
          f"bawah={cfg['margin_bottom']}  "
          f"dalam={cfg['margin_dalam']}  "
          f"luar={cfg['margin_luar']}")
    print(f"  Mirror margin  : {'Ya' if cfg['mirror'] else 'Tidak'}")
    print(f"  Font isi       : {FONT_NAME} {cfg['font_body']}pt, "
          f"spasi {cfg['line_spacing']}")
    print()

    # Verifikasi file
    missing = [f for f in URUTAN_FILE if not f.exists()]
    if missing:
        print("⚠ File berikut tidak ditemukan:")
        for f in missing:
            print(f"    {f.relative_to(BASE_DIR)}")
        resp = input("\nLanjutkan tanpa file tersebut? (y/n): ").strip().lower()
        if resp != "y":
            print("Dibatalkan.")
            sys.exit(0)

    # Buat dokumen
    doc = setup_document(cfg)
    md = MarkdownParser(doc, cfg)

    total = len(URUTAN_FILE)
    for idx, filepath in enumerate(URUTAN_FILE, 1):
        filename = filepath.name
        is_cover = filename in COVER_FILES
        add_break = (idx > 1) and not is_cover

        rel = filepath.relative_to(BASE_DIR)
        print(f"  [{idx:2d}/{total}] {rel}")

        md.render_file(filepath, is_cover=is_cover, add_page_break=add_break)

    # Simpan
    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f"\n✅ Berhasil! File disimpan:")
    print(f"   {output_path}")
    print(f"   Ukuran: {size_kb:.0f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
