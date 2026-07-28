#!/usr/bin/env python3
"""
compile_panduan_bppu.py — Kompilasi panduan BPPU dari .md menjadi .docx

Dependensi:
    pip install python-docx

Contoh cepat:
    python compile_panduan_bppu.py                          # A4 default
    python compile_panduan_bppu.py --ukuran A5              # A5 untuk cetak kecil
    python compile_panduan_bppu.py --ukuran A4 --mirror     # A4 + mirror margin (jilid)
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx belum terinstal.")
    print("Jalankan:  pip install python-docx")
    sys.exit(1)

# -------------------------------------------------
# Konfigurasi path
# -------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
PANDUAN_DIR = BASE_DIR / "docs" / "panduan"

# Urutan file untuk kompilasi panduan
URUTAN_FILE = [
    PANDUAN_DIR / "00-daftar-isi.md",
    PANDUAN_DIR / "11-alur-utama.md",
    PANDUAN_DIR / "01-admin-officer.md",
    PANDUAN_DIR / "02-kantin.md",
    PANDUAN_DIR / "03-member-buyer.md",
    PANDUAN_DIR / "04a-praktik-baik-kantin.md",
    PANDUAN_DIR / "04b-praktik-baik-admin.md",
    PANDUAN_DIR / "04c-praktik-baik-mitra.md",
    PANDUAN_DIR / "05-mitra-usaha.md",
    PANDUAN_DIR / "06-pimpinan.md",
    PANDUAN_DIR / "08-sistem-poin.md",
    PANDUAN_DIR / "09-self-order.md",
    PANDUAN_DIR / "10-troubleshooting.md",
    # 07-arsitektur-teknis.md sengaja tidak diikutkan
    # karena terlalu teknis untuk pembaca panduan operasional
]

# Font utama
FONT_NAME = "Caladea"

# Warna branding BPPU IKIP Siliwangi
WARNA_EMAS = RGBColor(0x99, 0x66, 0x00)     # #996600
WARNA_HEADER_BG = "F4EFE5"                   # tint terang dari #996600

# -------------------------------------------------
# Preset ukuran halaman
# -------------------------------------------------
PAGE_PRESETS = {
    "A5": {
        "page_width": 14.8,
        "page_height": 21.0,
        "margin_top": 1.5,
        "margin_bottom": 1.5,
        "margin_dalam": 2.0,
        "margin_luar": 1.5,
        "font_body": 11,
        "font_h1": 14,
        "font_h2": 12,
        "font_h3": 11,
        "font_h4": 11,
        "font_table": 9,
        "font_code": 8,
        "font_caption": 9,
        "font_cover_title": 18,
        "font_cover_subtitle": 13,
        "font_cover_author": 12,
        "line_spacing": 1.15,
        "first_indent": 0,
        "list_indent": 0.8,
        "code_indent": 0.6,
    },
    "B5": {
        "page_width": 17.6,
        "page_height": 25.0,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_dalam": 2.5,
        "margin_luar": 1.8,
        "font_body": 11,
        "font_h1": 15,
        "font_h2": 13,
        "font_h3": 12,
        "font_h4": 11,
        "font_table": 9,
        "font_code": 8.5,
        "font_caption": 9,
        "font_cover_title": 20,
        "font_cover_subtitle": 14,
        "font_cover_author": 13,
        "line_spacing": 1.4,
        "first_indent": 0,
        "list_indent": 0.8,
        "code_indent": 0.8,
    },
    "A4": {
        "page_width": 21.0,
        "page_height": 29.7,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_dalam": 3.0,
        "margin_luar": 2.54,
        "font_body": 12,
        "font_h1": 16,
        "font_h2": 14,
        "font_h3": 13,
        "font_h4": 12,
        "font_table": 10,
        "font_code": 9,
        "font_caption": 10,
        "font_cover_title": 22,
        "font_cover_subtitle": 16,
        "font_cover_author": 14,
        "line_spacing": 1.5,
        "first_indent": 0,
        "list_indent": 1.0,
        "code_indent": 1.0,
    },
}


def get_config(args) -> dict:
    """Gabungkan preset dengan override CLI menjadi konfigurasi final."""
    preset_name = args.ukuran.upper()
    if preset_name in PAGE_PRESETS:
        cfg = dict(PAGE_PRESETS[preset_name])
    else:
        if not args.lebar or not args.tinggi:
            print(f"ERROR: Ukuran '{args.ukuran}' tidak dikenali.")
            print(f"Preset tersedia: {', '.join(PAGE_PRESETS.keys())}")
            print("Atau gunakan --lebar dan --tinggi untuk ukuran kustom.")
            sys.exit(1)
        cfg = dict(PAGE_PRESETS["A4"])
        cfg["page_width"] = args.lebar
        cfg["page_height"] = args.tinggi

    if args.margin_top is not None:
        cfg["margin_top"] = args.margin_top
    if args.margin_bottom is not None:
        cfg["margin_bottom"] = args.margin_bottom
    if args.margin_dalam is not None:
        cfg["margin_dalam"] = args.margin_dalam
    if args.margin_luar is not None:
        cfg["margin_luar"] = args.margin_luar
    if args.font_body is not None:
        cfg["font_body"] = args.font_body
    if args.spasi is not None:
        cfg["line_spacing"] = args.spasi

    cfg["mirror"] = args.mirror
    return cfg


# -------------------------------------------------
# Pemrosesan Markdown
# -------------------------------------------------

class MarkdownParser:
    """Parser Markdown sederhana yang langsung menulis ke python-docx."""

    def __init__(self, doc: Document, cfg: dict):
        self.doc = doc
        self.cfg = cfg

    @staticmethod
    def _clean_html_br(text: str) -> str:
        """Hapus tag <br> / <br/>."""
        return re.sub(r"<br\s*/?>", "\n", text)

    @staticmethod
    def _strip_links(text: str) -> str:
        """Ubah [teks](url) menjadi teks saja."""
        return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

    def _add_inline_formatting(self, paragraph, text: str):
        """Proses bold, italic, inline code, dan teks biasa ke runs."""
        text = self._clean_html_br(text)
        text = self._strip_links(text)

        pattern = re.compile(
            r"(\*\*\*(.+?)\*\*\*)"   # bold italic
            r"|(\*\*(.+?)\*\*)"      # bold
            r"|(\*(.+?)\*)"          # italic
            r"|(_(.+?)_)"            # italic alt
            r"|(`(.+?)`)"            # inline code
            r"|([^*_`]+)"            # teks biasa
        )

        for m in pattern.finditer(text):
            if m.group(2):  # bold italic
                run = paragraph.add_run(m.group(2))
                run.bold = True
                run.italic = True
            elif m.group(4):  # bold
                run = paragraph.add_run(m.group(4))
                run.bold = True
            elif m.group(6):  # italic
                run = paragraph.add_run(m.group(6))
                run.italic = True
            elif m.group(8):  # italic alt
                run = paragraph.add_run(m.group(8))
                run.italic = True
            elif m.group(10):  # inline code
                run = paragraph.add_run(m.group(10))
                run.font.name = "Consolas"
                run.font.size = Pt(self.cfg["font_code"])
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            elif m.group(11):  # teks biasa
                paragraph.add_run(m.group(11))

    def _format_paragraph(self, paragraph, font_size=None,
                          bold=False, italic=False,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=6,
                          first_line_indent=None):
        """Atur format paragraf."""
        if font_size is None:
            font_size = self.cfg["font_body"]

        pf = paragraph.paragraph_format
        pf.alignment = alignment
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = self.cfg["line_spacing"]

        if first_line_indent is not None:
            pf.first_line_indent = Cm(first_line_indent)

        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(font_size)
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'
                )
                rpr.append(rFonts)
            else:
                rFonts.set(qn("w:eastAsia"), FONT_NAME)

            if bold:
                run.bold = True
            if italic:
                run.italic = True

    # -- Tabel --

    def _parse_table(self, lines: list[str]) -> list[list[str]]:
        """Parse baris-baris tabel markdown menjadi list of rows."""
        rows = []
        for line in lines:
            line = line.strip()
            if not line.startswith("|"):
                continue
            if re.match(r"^\|[\s\-:|]+\|$", line):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            rows.append(cells)
        return rows

    def _add_table(self, rows: list[list[str]]):
        """Tambahkan tabel ke dokumen."""
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        ft = self.cfg["font_table"]
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    cell.paragraphs[0].clear()
                    self._add_inline_formatting(cell.paragraphs[0], cell_text)

                    for run in cell.paragraphs[0].runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(ft)
                        if i == 0:
                            run.bold = True

                    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.15

                    if i == 0:
                        shading = parse_xml(
                            f'<w:shd {nsdecls("w")} w:fill="{WARNA_HEADER_BG}" w:val="clear"/>'
                        )
                        cell._element.get_or_add_tcPr().append(shading)

        self.doc.add_paragraph()

    # -- Heading --

    def _add_heading(self, text: str, level: int):
        """Tambahkan heading ke dokumen."""
        size_map = {
            1: self.cfg["font_h1"],
            2: self.cfg["font_h2"],
            3: self.cfg["font_h3"],
            4: self.cfg["font_h4"],
        }
        font_size = size_map.get(level, self.cfg["font_body"])

        p = self.doc.add_paragraph()
        self._add_inline_formatting(p, text)
        alignment = (WD_ALIGN_PARAGRAPH.CENTER
                     if level <= 1
                     else WD_ALIGN_PARAGRAPH.LEFT)
        space_before = 24 if level == 1 else (18 if level == 2 else 12)
        space_after = 12 if level <= 2 else 6
        self._format_paragraph(p, font_size=font_size, bold=True,
                               alignment=alignment,
                               space_before=space_before,
                               space_after=space_after)

        # Garis bawah untuk H1
        if level == 1:
            for run in p.runs:
                run.font.color.rgb = WARNA_EMAS

    # -- List --

    def _add_list_item(self, text: str, ordered: bool = False,
                       number: int = 1, indent_level: int = 0):
        """Tambahkan item daftar."""
        p = self.doc.add_paragraph()
        prefix = f"{number}. " if ordered else "- "
        p.add_run(prefix)
        self._add_inline_formatting(p, text)
        self._format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               space_before=0, space_after=3,
                               first_line_indent=0)
        indent = self.cfg["list_indent"] + (indent_level * 0.5)
        p.paragraph_format.left_indent = Cm(indent)

    # -- Blok kode --

    def _add_code_block(self, lines: list[str]):
        """Tambahkan blok kode."""
        for line in lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(self.cfg["font_code"])
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Cm(self.cfg["code_indent"])

    # -- Blockquote / Catatan --

    def _add_blockquote(self, text: str):
        """Tambahkan blok kutipan/catatan."""
        # Hapus marker [!NOTE], [!TIP], [!IMPORTANT], [!WARNING], [!CAUTION]
        clean = re.sub(r"\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]\s*", "", text)
        clean = clean.strip()
        if not clean:
            return

        p = self.doc.add_paragraph()
        run = p.add_run("Catatan: ")
        run.bold = True
        self._add_inline_formatting(p, clean)
        self._format_paragraph(p, italic=True,
                               alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                               space_before=6, space_after=6,
                               first_line_indent=0)
        p.paragraph_format.left_indent = Cm(0.5)

    # -- Halaman sampul --

    def _render_cover(self):
        """Render halaman sampul BPPU."""
        blank = 4 if self.cfg["page_height"] <= 21.5 else 6
        for _ in range(blank):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)

        # Judul
        p = self.doc.add_paragraph()
        run = p.add_run("PANDUAN SISTEM")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(self.cfg["font_cover_title"])
        run.font.color.rgb = WARNA_EMAS
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        p = self.doc.add_paragraph()
        run = p.add_run("BPPU IKIP SILIWANGI")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(self.cfg["font_cover_title"])
        run.font.color.rgb = WARNA_EMAS
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)

        # Subtitle
        p = self.doc.add_paragraph()
        run = p.add_run("Badan Pengelola dan Pengembangan Usaha")
        run.italic = True
        run.font.name = FONT_NAME
        run.font.size = Pt(self.cfg["font_cover_subtitle"])
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)

        p = self.doc.add_paragraph()
        run = p.add_run("Panduan Operasional Lengkap untuk Seluruh Pengguna")
        run.italic = True
        run.font.name = FONT_NAME
        run.font.size = Pt(self.cfg["font_cover_subtitle"])
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(40)

        # Info bawah
        info_lines = [
            "https://bppu.ikipsiliwangi.ac.id",
            "bppu@ikipsiliwangi.ac.id",
            "",
            "Versi 2.0 -- Juli 2026",
        ]
        for line in info_lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(self.cfg["font_cover_author"])
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)

    # -- Garis pemisah (---) diubah menjadi Page Break --

    def _add_page_break(self):
        """Tambahkan page break."""
        self.doc.add_page_break()

    # -- Parser utama per file --

    def render_file(self, filepath: Path, add_page_break: bool = True):
        """Parse dan render satu file markdown ke dokumen."""
        if not filepath.exists():
            print(f"  [!] File tidak ditemukan: {filepath}")
            return

        text = filepath.read_text(encoding="utf-8")
        lines = text.split("\n")

        if add_page_break:
            self.doc.add_page_break()

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Baris kosong
            if not stripped:
                i += 1
                continue

            # Garis pemisah ---
            if re.match(r"^-{3,}$", stripped):
                self._add_page_break()
                i += 1
                continue

            # Metadata (Versi:, Terakhir diperbarui:, dll) - render sebagai teks kecil
            if stripped.startswith("**Versi:") or stripped.startswith("**Terakhir"):
                p = self.doc.add_paragraph()
                self._add_inline_formatting(p, stripped)
                self._format_paragraph(p, font_size=self.cfg["font_caption"],
                                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                       space_before=0, space_after=3)
                i += 1
                continue

            # Blok kode
            if stripped.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1
                self._add_code_block(code_lines)
                continue

            # Blockquote / catatan
            if stripped.startswith("> "):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith("> "):
                    quote_lines.append(lines[i].strip()[2:])
                    i += 1
                self._add_blockquote(" ".join(quote_lines))
                continue

            # Tabel
            if stripped.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                rows = self._parse_table(table_lines)
                self._add_table(rows)
                continue

            # Heading
            heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                self._add_heading(heading_match.group(2), level)
                i += 1
                continue

            # Sub-list (indented)
            sub_match = re.match(r"^\s{2,}[-*]\s+(.+)$", line)
            if sub_match:
                self._add_list_item(sub_match.group(1), ordered=False,
                                    indent_level=1)
                i += 1
                continue

            sub_ol_match = re.match(r"^\s{2,}(\d+)\.\s+(.+)$", line)
            if sub_ol_match:
                num = int(sub_ol_match.group(1))
                self._add_list_item(sub_ol_match.group(2), ordered=True,
                                    number=num, indent_level=1)
                i += 1
                continue

            # Ordered list
            ol_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
            if ol_match:
                num = int(ol_match.group(1))
                self._add_list_item(ol_match.group(2), ordered=True,
                                    number=num)
                i += 1
                continue

            # Unordered list
            if stripped.startswith("- ") or stripped.startswith("* "):
                self._add_list_item(stripped[2:], ordered=False)
                i += 1
                continue

            # Paragraf biasa
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                nl = lines[i].strip()
                if (not nl
                        or nl.startswith("#")
                        or nl.startswith("|")
                        or nl.startswith("```")
                        or nl.startswith("> ")
                        or nl.startswith("- ")
                        or nl.startswith("* ")
                        or re.match(r"^\d+\.\s+", nl)
                        or re.match(r"^-{3,}$", nl)):
                    break
                para_lines.append(nl)
                i += 1

            full_text = " ".join(para_lines)
            p = self.doc.add_paragraph()
            self._add_inline_formatting(p, full_text)
            self._format_paragraph(
                p,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                space_before=0, space_after=6,
                first_line_indent=self.cfg["first_indent"],
            )


# -------------------------------------------------
# Setup dokumen
# -------------------------------------------------

def setup_document(cfg: dict) -> Document:
    """Buat dokumen baru dengan pengaturan halaman sesuai konfigurasi."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(cfg["page_width"])
    section.page_height = Cm(cfg["page_height"])
    section.top_margin = Cm(cfg["margin_top"])
    section.bottom_margin = Cm(cfg["margin_bottom"])

    if cfg["mirror"]:
        section.left_margin = Cm(cfg["margin_dalam"])
        section.right_margin = Cm(cfg["margin_luar"])
        sectPr = section._sectPr
        mirrorMargins = sectPr.find(qn("w:mirrorMargins"))
        if mirrorMargins is None:
            sectPr.append(parse_xml(
                f'<w:mirrorMargins {nsdecls("w")}/>'
            ))
    else:
        section.left_margin = Cm(cfg["margin_dalam"])
        section.right_margin = Cm(cfg["margin_luar"])

    # Style default
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(cfg["font_body"])
    style.paragraph_format.line_spacing = cfg["line_spacing"]
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'
        )
        rpr.append(rFonts)

    return doc


# -------------------------------------------------
# Main
# -------------------------------------------------

def build_parser() -> argparse.ArgumentParser:
    """Bangun argument parser dengan semua opsi."""
    p = argparse.ArgumentParser(
        description="Kompilasi Panduan BPPU dari .md menjadi .docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
contoh penggunaan:
  python compile_panduan_bppu.py                                # A4 default
  python compile_panduan_bppu.py --ukuran A5                    # A5 cetak kecil
  python compile_panduan_bppu.py --ukuran A4 --mirror           # A4 + mirror margin
  python compile_panduan_bppu.py --ukuran A5 --margin-dalam 2.5 # override margin
  python compile_panduan_bppu.py --font-body 11 --spasi 1.3     # override tipografi
        """,
    )

    p.add_argument(
        "--output", "-o",
        default="Panduan_Sistem_BPPU.docx",
        help="Nama file output (default: Panduan_Sistem_BPPU.docx)",
    )

    g_page = p.add_argument_group("ukuran halaman")
    g_page.add_argument(
        "--ukuran", "-u",
        default="A4",
        help="Preset ukuran: A4 (default), A5, B5, "
             "atau 'custom' (wajib --lebar --tinggi)",
    )
    g_page.add_argument("--lebar", type=float, default=None,
                        help="Lebar halaman dalam cm (untuk ukuran custom)")
    g_page.add_argument("--tinggi", type=float, default=None,
                        help="Tinggi halaman dalam cm (untuk ukuran custom)")

    g_margin = p.add_argument_group("margin (cm)")
    g_margin.add_argument("--margin-top", type=float, default=None,
                          help="Margin atas")
    g_margin.add_argument("--margin-bottom", type=float, default=None,
                          help="Margin bawah")
    g_margin.add_argument("--margin-dalam", type=float, default=None,
                          help="Margin dalam / sisi jilid (gutter)")
    g_margin.add_argument("--margin-luar", type=float, default=None,
                          help="Margin luar / sisi potong")
    g_margin.add_argument("--mirror", action="store_true", default=False,
                          help="Aktifkan mirror margin (gutter bergantian "
                               "kiri-kanan untuk cetak bolak-balik)")

    g_typo = p.add_argument_group("tipografi")
    g_typo.add_argument("--font-body", type=float, default=None,
                        help="Ukuran font isi (pt)")
    g_typo.add_argument("--spasi", type=float, default=None,
                        help="Spasi baris (misal 1.3, 1.5)")

    return p


def main():
    parser = build_parser()
    args = parser.parse_args()
    cfg = get_config(args)

    output_path = BASE_DIR / args.output
    ukuran_label = args.ukuran.upper()

    print("=" * 60)
    print("  Kompilasi Panduan Sistem BPPU IKIP Siliwangi")
    print("=" * 60)

    print(f"\n  Ukuran halaman : {ukuran_label} "
          f"({cfg['page_width']}x{cfg['page_height']} cm)")
    print(f"  Margin (cm)    : atas={cfg['margin_top']}  "
          f"bawah={cfg['margin_bottom']}  "
          f"dalam={cfg['margin_dalam']}  "
          f"luar={cfg['margin_luar']}")
    print(f"  Mirror margin  : {'Ya' if cfg['mirror'] else 'Tidak'}")
    print(f"  Font isi       : {FONT_NAME} {cfg['font_body']}pt, "
          f"spasi {cfg['line_spacing']}")
    print()

    # Verifikasi file
    missing = [f for f in URUTAN_FILE if not f.exists()]
    if missing:
        print("[!] File berikut tidak ditemukan:")
        for f in missing:
            print(f"    {f.relative_to(BASE_DIR)}")
        resp = input("\nLanjutkan tanpa file tersebut? (y/n): ").strip().lower()
        if resp != "y":
            print("Dibatalkan.")
            sys.exit(0)

    # Buat dokumen
    doc = setup_document(cfg)
    md = MarkdownParser(doc, cfg)

    # Render cover
    print("  [ 0] Halaman Sampul")
    md._render_cover()

    # Render semua file
    total = len(URUTAN_FILE)
    for idx, filepath in enumerate(URUTAN_FILE, 1):
        rel = filepath.relative_to(BASE_DIR)
        print(f"  [{idx:2d}/{total}] {rel}")
        md.render_file(filepath, add_page_break=True)

    # Simpan
    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f"\nBerhasil! File disimpan:")
    print(f"   {output_path}")
    print(f"   Ukuran: {size_kb:.0f} KB")
    print("=" * 60)


if __name__ == "__main__":
    main()
