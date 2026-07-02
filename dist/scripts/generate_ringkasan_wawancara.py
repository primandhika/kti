from __future__ import annotations

import csv
from collections import Counter, defaultdict
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "data" / "kualitatif" / "wawancara.csv"
OUT_DIR = ROOT / "outputs"
OUT_THEME = OUT_DIR / "ringkasan_hasil_wawancara_tematik.csv"
OUT_DETAIL = OUT_DIR / "ringkasan_hasil_wawancara_detail.csv"
OUT_DOCX = OUT_DIR / "ringkasan_hasil_wawancara.docx"


THEME_SUMMARIES = {
    "Peningkatan Kesadaran Metakognitif": (
        "Wawancara menunjukkan bahwa media membantu mahasiswa menyadari perlunya strategi, "
        "perencanaan, pemantauan tempo/intonasi, dan identifikasi kelemahan sebelum atau selama "
        "berbicara. Dosen juga melihat sebagian mahasiswa mulai memahami bahwa kegiatan berbicara "
        "memerlukan persiapan."
    ),
    "Kemudahan Akses dan Fleksibilitas": (
        "Mahasiswa menilai media mudah diakses secara mandiri, dapat diputar ulang, dan memiliki "
        "durasi konten yang relatif sesuai. Fleksibilitas ini mendorong keterlibatan belajar di luar "
        "jadwal kelas."
    ),
    "Peningkatan Kepercayaan Diri": (
        "Latihan berulang dan penguasaan materi melalui media berkontribusi pada rasa lebih tenang, "
        "berkurangnya kecemasan berbicara, dan meningkatnya keyakinan ketika menyampaikan materi."
    ),
    "Pengorganisasian Ide Terstruktur": (
        "Media membantu mahasiswa menyusun kerangka sebelum berbicara, terutama pada pembukaan, isi, "
        "dan penutup. Namun, terdapat mahasiswa yang masih kesulitan mentransfer konsep urutan "
        "penyampaian ke praktik nyata."
    ),
    "Refleksi dan Evaluasi Diri": (
        "Fitur playback, rekaman, umpan balik, dan latihan berulang mendorong mahasiswa menilai "
        "kembali performa berbicaranya. Proses ini tampak pada kesadaran kesalahan, revisi bertahap, "
        "dan perbaikan cara menyampaikan pesan."
    ),
    "Keterlibatan Mahasiswa": (
        "Dosen melihat keterlibatan mahasiswa kelas eksperimen cukup baik, meskipun kendala perangkat "
        "dan login membuat sebagian progres tidak tercatat. Pada kelas kontrol, penggunaan media "
        "bersifat sukarela dan hanya dilakukan oleh sebagian kecil mahasiswa."
    ),
    "Efektivitas Media": (
        "Media dipandang efektif sebagai pendamping pembelajaran karena menyediakan akses rutin, latihan "
        "antarseksi, dan dorongan penyelesaian melalui poin. Efeknya paling terlihat pada mahasiswa yang "
        "menggunakan media secara konsisten."
    ),
    "Fleksibilitas Penggunaan": (
        "Fitur rekaman memberi ruang latihan, tetapi menuntut kondisi lingkungan yang sepi agar hasil "
        "rekaman jelas. Sebagian mahasiswa kemudian melanjutkan penggunaan media dari rumah."
    ),
    "Tantangan Penerapan": (
        "Kendala penerapan mencakup adaptasi dosen terhadap fitur penilaian digital, beban menjawab "
        "pertanyaan mahasiswa, akses fitur dari ponsel, serta kebutuhan alur kerja yang lebih ringkas."
    ),
    "Saran Pengembangan": (
        "Saran utama dosen mengarah pada penilaian otomatis berbantuan AI, mode offline, kemudahan akses "
        "mobile, dan fitur simpan otomatis untuk mengurangi beban teknis."
    ),
    "Karakteristik Belajar Mahasiswa": (
        "Kelas kontrol cenderung menggunakan sumber bebas, bergantung pada gawai, dan memiliki minat "
        "baca yang rendah. Kondisi ini membuat kualitas persiapan belajar tidak seragam."
    ),
    "Perbandingan Pembelajaran": (
        "Pembelajaran kelas kontrol berlangsung dengan PBL reguler dan sumber belajar yang lebih bebas, "
        "sedangkan kelas eksperimen memiliki tuntutan penyelesaian modul yang lebih terarah."
    ),
    "Keterampilan Berbicara": (
        "Dosen menyoroti aspek kebahasaan dan kesantunan sebagai bagian yang masih perlu dibenahi, "
        "terutama penggunaan ragam bahasa formal dalam situasi akademik."
    ),
}


def clean_text(value: object) -> str:
    if pd.isna(value):
        return ""
    text = str(value)
    replacements = {
        "â€”": "-",
        "â€“": "-",
        "â€˜": "'",
        "â€™": "'",
        "â€œ": '"',
        "â€\x9d": '"',
        "\u00a0": " ",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return " ".join(text.split())


def source_label(code: str) -> str:
    if code.startswith("D"):
        return "Dosen"
    if code.startswith("E"):
        return "Mahasiswa"
    return "Lainnya"


def load_data() -> pd.DataFrame:
    df = pd.read_csv(SOURCE, dtype=str, keep_default_na=False)
    for column in df.columns:
        df[column] = df[column].map(clean_text)
    df["jenis_sumber"] = df["kode_subjek"].map(source_label)
    return df


def summarize_themes(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for number, (theme, group) in enumerate(df.groupby("tema", sort=False), start=1):
        subthemes = "; ".join(dict.fromkeys(group["subtema"].tolist()))
        source_counts = Counter(group["jenis_sumber"])
        source = ", ".join(f"{name} ({count})" for name, count in sorted(source_counts.items()))
        dates = [value for value in dict.fromkeys(group["tanggal_wawancara"].tolist()) if value]
        rows.append(
            {
                "no": number,
                "tema": theme,
                "jumlah_data": len(group),
                "subtema": subthemes,
                "sumber_data": source,
                "tanggal_wawancara": "; ".join(dates),
                "ringkasan_temuan": THEME_SUMMARIES.get(theme, ""),
            }
        )
    return pd.DataFrame(rows)


def detail_table(df: pd.DataFrame) -> pd.DataFrame:
    detail = df[
        [
            "no",
            "jenis_sumber",
            "kode_subjek",
            "nama",
            "tema",
            "subtema",
            "kutipan",
            "catatan_analisis",
            "tanggal_wawancara",
        ]
    ].copy()
    detail = detail.rename(
        columns={
            "jenis_sumber": "sumber",
            "kutipan": "kutipan_kunci",
            "catatan_analisis": "makna_analisis",
        }
    )
    return detail


def set_document_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_table(document: Document, columns: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for index, column in enumerate(columns):
        header_cells[index].text = column
        for paragraph in header_cells[index].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    document.add_paragraph()


def build_docx(df: pd.DataFrame, themes: pd.DataFrame, detail: pd.DataFrame) -> None:
    document = Document()
    set_document_style(document)
    add_heading(document, "RINGKASAN HASIL WAWANCARA")

    subject_counts = defaultdict(set)
    for _, row in df.iterrows():
        subject_counts[row["jenis_sumber"]].add(row["kode_subjek"])
    mahasiswa_n = len(subject_counts["Mahasiswa"])
    dosen_n = len(subject_counts["Dosen"])
    dates = [value for value in dict.fromkeys(df["tanggal_wawancara"].tolist()) if value]
    add_paragraph(
        document,
        "Ringkasan ini disusun dari data wawancara mahasiswa dan dosen pada tahap implementasi "
        f"media Bicaranta. Data mencakup {len(df)} satuan kutipan/temuan, terdiri atas "
        f"{mahasiswa_n} mahasiswa dan {dosen_n} dosen. Tanggal wawancara yang tercatat: "
        f"{'; '.join(dates)}.",
    )

    add_paragraph(document, "A. Rekapitulasi Tematik")
    add_table(
        document,
        ["No", "Tema", "Jml.", "Subtema", "Sumber Data", "Ringkasan Temuan"],
        [
            [
                row["no"],
                row["tema"],
                row["jumlah_data"],
                row["subtema"],
                row["sumber_data"],
                row["ringkasan_temuan"],
            ]
            for _, row in themes.iterrows()
        ],
    )

    add_paragraph(document, "B. Matriks Kutipan Kunci dan Makna Analisis")
    add_table(
        document,
        ["No", "Sumber", "Kode", "Tema/Subtema", "Kutipan Kunci", "Makna Analisis"],
        [
            [
                row["no"],
                row["sumber"],
                row["kode_subjek"],
                f"{row['tema']} / {row['subtema']}",
                row["kutipan_kunci"],
                row["makna_analisis"],
            ]
            for _, row in detail.iterrows()
        ],
    )

    document.save(OUT_DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    df = load_data()
    themes = summarize_themes(df)
    detail = detail_table(df)
    themes.to_csv(OUT_THEME, index=False, quoting=csv.QUOTE_MINIMAL, encoding="utf-8-sig")
    detail.to_csv(OUT_DETAIL, index=False, quoting=csv.QUOTE_MINIMAL, encoding="utf-8-sig")
    build_docx(df, themes, detail)
    print(f"Wrote {OUT_THEME.relative_to(ROOT)}")
    print(f"Wrote {OUT_DETAIL.relative_to(ROOT)}")
    print(f"Wrote {OUT_DOCX.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
